#!/usr/bin/env python3
"""
Summary-Only Regeneration Script
Regenerates summaries from existing transcription files without re-transcribing
"""

import os
import sys
import yaml
import logging
import argparse
from pathlib import Path
from typing import List, Optional, Dict
import requests
import datetime
from tqdm import tqdm

# Import the TranscriptionProcessor to reuse summary generation logic
from transcribe import TranscriptionProcessor


class SummaryRegenerator:
    def __init__(self, config_path: str = "config.yaml"):
        """Initialize with configuration"""
        with open(config_path, 'r') as f:
            self.config = yaml.safe_load(f)
        
        # Setup logging
        log_level = getattr(logging, self.config['processing']['log_level'])
        logging.basicConfig(
            level=log_level,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Setup paths
        self.input_folder = Path(self.config['folders']['input'])
        self.output_folder = Path(self.config['folders']['output']) if self.config['folders']['output'] else None
        
        # Check Ollama availability
        self.ollama_available = self._check_ollama()
        if not self.ollama_available:
            raise RuntimeError("Ollama is required for summary generation but not available")
        
        # Get summary regeneration config
        self.regen_config = self.config.get('summary_regeneration', {})
        if not self.regen_config.get('enabled', True):
            raise RuntimeError("Summary regeneration is disabled in config")

    def _check_ollama(self) -> bool:
        """Check if Ollama is available and model exists"""
        try:
            api_url = self.config['ollama']['api_url']
            response = requests.get(f"{api_url}/api/tags", timeout=2)
            if response.status_code == 200:
                models = response.json().get('models', [])
                model_names = [m['name'] for m in models]
                target_model = self.config['ollama']['model']
                
                model_found = any(target_model in name for name in model_names)
                
                if model_found:
                    self.logger.info(f"Ollama available with model: {target_model}")
                    return True
                else:
                    self.logger.error(f"Ollama model '{target_model}' not found.")
                    self.logger.info(f"Available models: {', '.join(model_names)}")
                    return False
            return False
        except Exception as e:
            self.logger.error(f"Ollama not available: {e}")
            return False

    def get_output_path(self, input_path: Path) -> Path:
        """Determine output path based on configuration (same logic as transcribe.py)"""
        if self.output_folder:
            self.output_folder.mkdir(parents=True, exist_ok=True)
            
            preserve_structure = self.config.get('output', {}).get('preserve_structure', True)
            
            if preserve_structure:
                try:
                    rel_path = input_path.relative_to(self.input_folder)
                    output_path = self.output_folder / rel_path.parent / input_path.name
                    output_path.parent.mkdir(parents=True, exist_ok=True)
                except ValueError:
                    output_path = self.output_folder / input_path.name
            else:
                output_path = self.output_folder / input_path.name
        else:
            output_path = input_path
        
        return output_path

    def find_transcription_files(self) -> List[Path]:
        """Find all .txt transcription files"""
        txt_files = []
        
        if self.config['processing']['recursive']:
            pattern = "**/*.txt"
        else:
            pattern = "*.txt"
        
        for txt_file in self.input_folder.glob(pattern):
            if txt_file.is_file():
                txt_files.append(txt_file)
        
        self.logger.info(f"Found {len(txt_files)} transcription files")
        return txt_files

    def should_skip_summary(self, output_path: Path, lang_code: str, force: bool = False) -> bool:
        """Check if summary should be skipped"""
        if force:
            return False
        
        if lang_code in ("no", "primary"):
            summary_path = output_path.with_suffix('.docx')
        else:
            # FIX: Construct path correctly
            summary_path = output_path.parent / f"{output_path.stem}_{lang_code}.docx"
        
        return summary_path.exists()

def generate_summary_for_language(self, transcription_text: str, output_path: Path,
                                  lang_code: str, detected_lang: str = "unknown") -> bool:
    """Generate summary in specified language"""
    try:
        max_length = self.config["ollama"]["summary"]["max_length"]
        style = self.config["ollama"]["summary"]["style"]
        extract_topics = self.config["ollama"]["summary"]["extract_topics"]

        system_prompt = ""
        prompt_prefix = ""

        if lang_code == "no":
            lang_instruction = "på norsk (bokmål)"
            lang_name = "Norwegian"

            system_prompt = (
                "Du er en norsk AI-assistent. Du må ALLTID svare på norsk.\n"
                "VIKTIG: Hele sammendraget skal skrives på norsk. Ikke bruk engelsk."
            )

            prompt_prefix = (
                "VIKTIG: Skriv hele sammendraget på NORSK (bokmål). "
                "Ikke bruk engelsk.\n\n"
            )
        elif lang_code == "en":
            lang_instruction = "in English"
            lang_name = "English"
            system_prompt = "You are an English-speaking AI assistant. Write only in English."
        else:
            lang_instruction = f"in {lang_code.upper()}"
            lang_name = lang_code.upper()

        if style == "concise":
            if lang_code == "no":
                style_instruction = "Skriv et kort sammendrag på norsk (2-3 setninger)"
            else:
                style_instruction = "Write a concise summary (2-3 sentences)"
        elif style == "bullet_points":
            if lang_code == "no":
                style_instruction = "Skriv et sammendrag på norsk som punktliste med viktige punkter"
            else:
                style_instruction = "Write a summary as bullet points highlighting key information"
        else:
            if lang_code == "no":
                style_instruction = (
                    "Skriv et detaljert sammendrag på norsk som dekker alle "
                    "hovedpunkter og viktig informasjon"
                )
            else:
                style_instruction = (
                    "Write a detailed summary covering all main points and key information"
                )

        if lang_code == "no":
            prompt = f"""{system_prompt}

{prompt_prefix}{style_instruction} {lang_instruction}.

Transkripsjon (oppdaget språk: {detected_lang.upper()}):

{transcription_text}

Sammendrag på norsk:"""
        else:
            prompt = f"""{system_prompt}

{style_instruction} {lang_instruction}.

Transcription (detected language: {detected_lang.upper()}):

{transcription_text}

Summary:"""

        if extract_topics:
            if lang_code == "no":
                prompt += "\n\n[Etter sammendraget, list 3-5 nøkkeltemaer på norsk]"
            else:
                prompt += f"\n\n[After the summary, list 3-5 key topics/themes in {lang_name}]"

        api_url = self.config["ollama"]["api_url"]
        model = self.config["ollama"]["model"]

        response = requests.post(
            f"{api_url}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.5 if lang_code == "no" else 0.7,
                    "num_predict": max_length * 2,
                },
            },
            timeout=120,
        )

        if response.status_code != 200:
            self.logger.error(f"Ollama API error for {lang_name}: {response.status_code}")
            return False

        summary_text = response.json()["response"]

        if lang_code in ("no", "primary"):
            docx_path = output_path.with_suffix(".docx")
            md_path = output_path.with_suffix(".md")
        else:
            docx_path = output_path.parent / f"{output_path.stem}_{lang_code}.docx"
            md_path = output_path.parent / f"{output_path.stem}_{lang_code}.md"

        # skriv en ren tekst eller markdown versjon for HTML
        try:
            with open(md_path, "w", encoding="utf-8") as f_md:
                f_md.write(summary_text.strip())
            self.logger.info(f"Saved plain summary for HTML: {md_path}")
        except Exception as e:
            self.logger.warning(f"Could not write markdown summary: {e}")

        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            doc = Document()

            title_text = (
                f"Sammendrag: {output_path.stem}"
                if lang_code == "no"
                else f"Summary: {output_path.stem}"
            )
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            meta = doc.add_paragraph()
            meta.add_run(f"Generert av: {model}\n").bold = True
            meta.add_run(f"Språk: {lang_name}\n").bold = True
            meta.add_run(f"Kildespråk: {detected_lang.upper()}\n").bold = True

            doc.add_paragraph("_" * 50)

            for line in summary_text.strip().split("\n"):
                line = line.strip()
                if not line:
                    continue

                if line.startswith("#"):
                    heading_text = line.lstrip("#").strip()
                    doc.add_heading(heading_text, level=2)
                elif line.isupper() and len(line) < 50:
                    doc.add_heading(line, level=2)
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    p.add_run(line.strip("*")).bold = True
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line[0].isdigit() and ". " in line[:4]:
                    doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
                else:
                    doc.add_paragraph(line)

            doc.add_paragraph("_" * 50)
            footer = doc.add_paragraph()
            if lang_code == "no":
                footer_text = f"Generert fra transkripsjon: {output_path.with_suffix('.txt').name}"
            else:
                footer_text = f"Generated from transcription: {output_path.with_suffix('.txt').name}"
            footer.add_run(footer_text).italic = True

            doc.save(str(docx_path))
            self.logger.info(f"Generated {lang_name} summary: {docx_path}")
            return True

        except ImportError:
            self.logger.warning("python-docx not available, using only markdown")
            return True

    except Exception as e:
        self.logger.error(f"Summary generation failed for {lang_code}: {e}")
        return False

    def process_transcription_file(self, txt_file: Path, target_languages: List[str] = None, 
                                 force: bool = False):
        """Process a single transcription file"""
        self.logger.info(f"Processing: {txt_file.name}")
        
        # Read transcription
        try:
            with open(txt_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extract just the text (remove timestamps if present)
            lines = content.split('\n')
            text_lines = []
            detected_lang = "unknown"
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Skip metadata lines
                if line.startswith('Language:') or line.startswith('Detected language:'):
                    if 'Language:' in line:
                        detected_lang = line.split(':', 1)[1].strip()
                    continue
                
                # Remove timestamps [HH:MM:SS] or [123.45s]
                import re
                line = re.sub(r'\[\d{2}:\d{2}:\d{2}\]', '', line)
                line = re.sub(r'\[\d+\.?\d*s?\]', '', line)
                line = line.strip()
                
                if line:
                    text_lines.append(line)
            
            transcription_text = ' '.join(text_lines)
            
            if len(transcription_text.split()) < 50:
                self.logger.info(f"Transcription too short for summary: {txt_file.name}")
                return
            
        except Exception as e:
            self.logger.error(f"Failed to read transcription {txt_file.name}: {e}")
            return
        
        # Get output path
        output_path = self.get_output_path(txt_file)
        
        # Determine languages to generate
        if target_languages:
            languages = target_languages
        else:
            languages = self.regen_config.get('languages', ['no', 'en'])
        
        # Generate summaries for each language
        for lang_code in languages:
            if self.should_skip_summary(output_path, lang_code, force):
                self.logger.info(f"Skipping existing {lang_code} summary for {txt_file.name}")
                continue
            
            success = self.generate_summary_for_language(
                transcription_text, output_path, lang_code, detected_lang
            )
            
            if not success:
                self.logger.warning(f"Failed to generate {lang_code} summary for {txt_file.name}")

    def process_all(self, target_languages: List[str] = None, force: bool = False):
        """Process all transcription files"""
        txt_files = self.find_transcription_files()
        
        if not txt_files:
            self.logger.info("No transcription files found")
            return
        
        self.logger.info(f"Processing {len(txt_files)} transcription files...")
        
        for txt_file in tqdm(txt_files, desc="Generating summaries"):
            self.process_transcription_file(txt_file, target_languages, force)
        
        self.logger.info("Summary regeneration complete!")


def main():
    """Main entry point with command line arguments"""
    parser = argparse.ArgumentParser(description="Regenerate summaries from existing transcriptions")
    parser.add_argument('--config', default='config.yaml', help='Configuration file path')
    parser.add_argument('--force', action='store_true', help='Overwrite existing summaries')
    parser.add_argument('--lang', action='append', help='Target language(s) (can be used multiple times)')
    parser.add_argument('--file', help='Process specific transcription file')
    
    args = parser.parse_args()
    
    try:
        regenerator = SummaryRegenerator(args.config)
        
        if args.file:
            # Process single file
            txt_file = Path(args.file)
            if not txt_file.exists():
                logging.error(f"File not found: {txt_file}")
                sys.exit(1)
            
            regenerator.process_transcription_file(txt_file, args.lang, args.force)
        else:
            # Process all files
            regenerator.process_all(args.lang, args.force)
            
    except Exception as e:
        logging.error(f"Fatal error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
